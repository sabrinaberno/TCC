import pandas as pd
from scipy.stats import spearmanr
import matplotlib.pyplot as plt
from docx import Document
import os

file_path = "Planilha_normalizada.xlsx"
data = pd.read_excel(file_path)

numerical_columns = [
    "Idade", "Estágio", "Morfo", "KIDScore", "t2", "t3", "t4", "t5", "t8", "tSC", "tSB", "tB", "cc2 (t3-t2)", 
    "cc3 (t5-t3)", "t5-t2", "s2 (t4-t3)", "s3 (t8-t5)", "tSC-t8", "tB-tSB", "Ploidia"
]

numerical_columns = [col for col in numerical_columns if col in data.columns]

generated_pairs = set()
results = []

for col1 in numerical_columns:
    for col2 in numerical_columns:
        if col1 != col2:  # Não exclui mais a correlação inversa
            coef, _ = spearmanr(data[col1], data[col2], nan_policy="omit")
            results.append({"Variable 1": col1, "Variable 2": col2, "Spearman Coefficient": coef})

correlation_df = pd.DataFrame(results)
output_file = "correlation_results.xlsx"
correlation_df.to_excel(output_file, index=False)
print(f"\nResultados salvos no arquivo: {output_file}")

correlation_df = pd.read_excel(output_file)

output_dir = "graficos"
os.makedirs(output_dir, exist_ok=True)

doc = Document()
doc.add_heading('Correlação de Spearman - Embriões', 0)

color_var1 = "lightgreen"
color_var2 = "darkblue"

for index, row in correlation_df.iterrows():
    var1, var2 = row['Variable 1'], row['Variable 2']
    coef = row['Spearman Coefficient']

    if not pd.isnull(coef):
        plt.figure(figsize=(6, 4))

        plt.scatter(data[var1], data[var2], alpha=0.6, c=color_var1, label=f'{var1} vs {var2}', marker='o')
        plt.scatter(data[var2], data[var1], alpha=0.6, c=color_var2, label=f'{var2}', marker='x')
        plt.xlabel(var1)
        plt.ylabel(var2)
        plt.legend(title="Variáveis")

        sanitized_var1 = var1.replace(" ", "_").replace("/", "-")
        sanitized_var2 = var2.replace(" ", "_").replace("/", "-")
        pdf_filename = f"grafico_{sanitized_var1}_vs_{sanitized_var2}_Spearman_{coef:.2f}.pdf"
        pdf_path = os.path.join(output_dir, pdf_filename)
        plt.savefig(pdf_path, format="pdf")
        plt.close()

        doc.add_heading(f'Correlação entre {var1} e {var2}', level=1)
        doc.add_paragraph(f'Coeficiente de Spearman: {coef:.2f}')
        doc.add_paragraph(f'O gráfico correspondente foi salvo como: {pdf_filename}')

output_word = "relatorio_correlacoes.docx"
doc.save(output_word)

print(f"\nRelatório gerado e salvo em: {output_word}")